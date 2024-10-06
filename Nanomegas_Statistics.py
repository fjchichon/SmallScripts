# -*- coding: utf-8 -*-
"""
Created on Mon Sep 30 23:38:29 2024

@author: Conway
"""

# %%%%%%%% STADISTICOS ·····plot···· con varios parametros de entrada
######################################


# -*- coding: utf-8 -*-
"""
Created on Fri Sep 20 19:19:25 2024

@author: Conway
"""
import os
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from scipy import stats
import scikit_posthocs as sp
from docx import Document
from docx.table import Table, _Cell
from scipy.stats import kruskal
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from scipy.cluster.hierarchy import dendrogram, linkage



## COLORS 
sns.set(style="whitegrid")
plt.ion()
# Create an array with the colors you want to use

# Crear paleta de colores azules, de más intenso a pastel
blues = sns.color_palette("Blues", 10)  # Paleta de 10 tonos de azul

# Crear paleta de colores naranjas, de más intenso a pastel
oranges = sns.color_palette("Oranges", 10)  # Paleta de 10 tonos de naranja

# Crear paleta nueva de 3 colores azules
#blues_palette = sns.color_palette("Blues", 3)

# %%%%%%%%% Lectura y generacion %%

#Leemos los ficheros csv
# Lista para almacenar los DataFrames
dataframes = []

# CSV Created in ImageJ
########################################################################################################################
# ,Area,X,Y,Perim.,Major,Minor,Angle,Circ.,Feret,Median,Skew,Kurt,FeretX,FeretY,FeretAngle,MinFeret,AR,Round,Solidity
#########################################################################################################################

# introducir numero de muestras ·····
ns = 6  # Cambia este valor según el número de muestras

# Bucle para cargar y añadir la columna 'Sample' a cada DataFrame
for nsample in range(1, ns + 1):
    # Leer el archivo CSV correspondiente
    file_name = f'S{nsample}_CLASS_1.csv'
    df = pd.read_csv(file_name)
    

    # Calcular la nueva columna AverageFeret como el promedio de 'Feret' y 'MinFeret'
    df['AvFeret'] = (df['Feret'] + df['MinFeret']) / 2
    
    # Calcular la nueva columna GD (Geometric Diameter) como 'Perim.' / π
    df['GD'] = df['Perim.'] / np.pi  # np.pi proporciona el valor de π
    
    # Calcular la nueva columna AED (Area Equivalent Diameter) como 2 * √(Area / π)
    df['AED'] = 2 * np.sqrt(df['Area'] / np.pi)
    
    # Añadir columna identificativa para cada muestra
    df['Sample'] = f'Sample {nsample}'
    
    # Añadir el DataFrame a la lista
    dataframes.append(df)
    
# Combinar todos los DataFrames en uno solo
df_combined = pd.concat(dataframes, ignore_index=True)

# Mostrar las primeras filas del DataFrame combinado para verificar
print(df_combined.head())





# %% ELIJE BIEN QUE PARAMETROS. Roundnesss y Corcularity irán en Naranja
# CSV Created in ImageJ
########################################################################################################################
# ,Area,X,Y,Perim.,Major,Minor,Angle,Circ.,Feret,Median,Skew,Kurt,FeretX,FeretY,FeretAngle,MinFeret,AR,Round,Solidity
#########################################################################################################################
    # Calcular la nueva columna AverageFeret (AVFeret) como el promedio de 'Feret' y 'MinFeret'
    # Calcular la nueva columna GD (Geometric Diameter) como 'Perim.' / π
    # Calcular la nueva columna AED (Area Equivalent Diameter) como 2 * √(Area / π)

# Te muestra pero no genera los plots:::::::::::::::::::::::::::::::::::::::::::::::::::::

# Definir diccionario de nombres largos
param_names = {
    'Feret': 'MaxFeret',
    'MinFeret': 'MinFeret',
    'Area': 'Area',
    'GD': 'Geometric Diameter (GD)',
    'Circ.': 'Circularity',
    'Round': 'Roundness',
    'AED': 'Area Equivalent Diameter (AED)',
    'Perim.': 'Perimeter',
    'AvFeret': 'Average Feret'
}

# Definir una lista de parámetros a analizar (nombres cortos usados en el DataFrame)
params = ['Feret', 'AvFeret', 'MinFeret', 'Area', 'GD', 'AED', 'Circ.', 'Round']  # Nombres en df_combined

# Iterar sobre cada parámetro en la lista
for param in params:
    long_name = param_names[param]  # Obtener el nombre largo

    print(f"\nAnálisis para el parámetro: {long_name}")

    # Realizar Kruskal-Wallis para el parámetro actual
    data_param = df_combined[['Sample', param]]  # Extraer solo la columna del parámetro actual
    all_data = [group[param].values for name, group in df_combined.groupby('Sample')]
    H_stat, p_value = kruskal(*all_data)
    print(f"Kruskal-Wallis H-Statistic: {H_stat}, p-value: {p_value}")

    if p_value < 0.05:
        print("Prueba de Kruskal-Wallis significativa. Realizando análisis post-hoc (Dunn's test)...")
        dunn_result = sp.posthoc_dunn(df_combined, val_col=param, group_col='Sample', p_adjust='bonferroni')
        print(dunn_result)
        
        # Crear heatmap para los resultados del Dunn's test
        plt.figure(figsize=(10, 8))
        sns.heatmap(dunn_result, annot=True, cmap='coolwarm', cbar=True, fmt=".2e")
        plt.title(f'Post-hoc Dunn’s Test Heatmap for {long_name}', fontsize=16)  # Aumentar tamaño de título
        #plt.savefig(f'Dunns_Test_Heatmap_{long_name}.png')
        plt.show()

    else:
        print("No se ha encontrado diferencia significativa.")

    # Elegir paleta de colores según el parámetro
    palette = oranges if param in ['Circ.', 'Round'] else blues

    # Crear el gráfico
    plt.figure(figsize=(14, 8))
    sns.boxplot(
    x='Sample', 
    y=param, 
    data=df_combined, 
    showfliers=False, 
    palette=palette,  # Aquí faltaba la coma
    showmeans=True, 
    linewidth=2, 
    meanprops={"marker":"X", "markerfacecolor":"black", "markeredgecolor":"white", "markersize":10}
    )

    # Ajustes del gráfico
    plt.title(f'Distribution of {long_name}', fontsize=20)  # Aumentar tamaño del título
    plt.xlabel('', fontsize=14)  # Aumentar tamaño del eje X
    ylabel_text = f'{long_name} (nm)' if param not in ['Area', 'Circ.', 'Round'] else (f'{long_name} (nm²)' if param == 'Area' else f'{long_name}')
    plt.ylabel(ylabel_text, fontsize=20)  # Aumentar tamaño del eje Y
    plt.xticks(fontsize=20)  # Aumentar tamaño de las etiquetas del eje X
    plt.yticks(fontsize=20)  # Aumentar tamaño de las etiquetas del eje Y
    
    # Guardar el gráfico en un archivo PNG con el nombre del parámetro
    #plt.savefig(f'{long_name}_Distribution.png')
    plt.show()

    print(f"Gráfico guardado como '{long_name}_Distribution.png'\n")
    
    
    
    
# %%%%%%%%% ALL, TEST, DENDOGRAMS and Boxplot %%


# Iterar sobre cada parámetro en la lista
for param in params:
    long_name = param_names[param]  # Obtener el nombre largo

    print(f"\nAnálisis para el parámetro: {long_name}")

    # Realizar Kruskal-Wallis para el parámetro actual
    data_param = df_combined[['Sample', param]]  # Extraer solo la columna del parámetro actual
    all_data = [group[param].values for name, group in df_combined.groupby('Sample')]
    H_stat, p_value = kruskal(*all_data)
    print(f"Kruskal-Wallis H-Statistic: {H_stat}, p-value: {p_value}")

    if p_value < 0.05:
        print("Prueba de Kruskal-Wallis significativa. Realizando análisis post-hoc (Dunn's test)...")
        dunn_result = sp.posthoc_dunn(df_combined, val_col=param, group_col='Sample', p_adjust='bonferroni')
        print(dunn_result)
        
        # Crear heatmap para los resultados del Dunn's test
        plt.figure(figsize=(10, 8))
        sns.heatmap(dunn_result, annot=True, cmap='coolwarm', cbar=True, fmt=".2e")
        plt.title(f'Post-hoc Dunn’s Test Heatmap for {long_name}', fontsize=16)  # Aumentar tamaño de título
        plt.savefig(f'Dunns_Test_Heatmap_{long_name}.png')
        plt.show()

        ### Clusterización jerárquica usando p-values transformados ###
        
        # Transformar p-values: 1 - p_value para convertir en similitud
        similarity_matrix = 1 - dunn_result.to_numpy()
        
        # Realizar la clusterización jerárquica (linkage)
        Z = linkage(similarity_matrix, method='ward')  # Puedes usar otros métodos como 'average', 'single', etc.
        
        # Crear dendrograma
        plt.figure(figsize=(10, 8))
        dendrogram(Z, labels=dunn_result.index, leaf_rotation=90, leaf_font_size=12)
        plt.title(f'Hierarchical Clustering Dendrogram for {long_name}', fontsize=16)
        #plt.xlabel('Muestras', fontsize=16)
        plt.ylabel('Distance', fontsize=16)
        plt.xticks(fontsize=20)  # Aumentar tamaño de las etiquetas del eje X
        plt.savefig(f'Dendrogram_{long_name}.png')
        plt.show()

    else:
        print("No se ha encontrado diferencia significativa.")

    # Gráficos de caja y bigotes (Boxplot)
    # Elegir paleta de colores según el parámetro
    palette = oranges if param in ['Circ.', 'Round'] else blues

    # Crear el gráfico
    plt.figure(figsize=(14, 8))
    sns.boxplot(
        x='Sample', 
        y=param, 
        data=df_combined, 
        showfliers=False, 
        palette=palette,  
        showmeans=True, 
        linewidth=2, 
        meanprops={"marker":"X", "markerfacecolor":"black", "markeredgecolor":"white", "markersize":10}
    )

    # Ajustes del gráfico
    plt.title(f'Distribution of {long_name}', fontsize=20)  # Aumentar tamaño del título
    plt.xlabel('', fontsize=14)  # Aumentar tamaño del eje X
    ylabel_text = f'{long_name} (nm)' if param not in ['Area', 'Circ.', 'Round'] else (f'{long_name} (nm²)' if param == 'Area' else f'{long_name}')
    plt.ylabel(ylabel_text, fontsize=20)  # Aumentar tamaño del eje Y
    plt.xticks(fontsize=20)  # Aumentar tamaño de las etiquetas del eje X
    plt.yticks(fontsize=20)  # Aumentar tamaño de las etiquetas del eje Y
    
    # Guardar el gráfico en un archivo PNG con el nombre del parámetro
    plt.savefig(f'{long_name}_Distribution.png')
    plt.show()

    print(f"Gráfico guardado como '{long_name}_Distribution.png'\n")



# %%%%%%%%% TABLES %%
#
# CSV Created in ImageJ
########################################################################################################################
# ,Area,X,Y,Perim.,Major,Minor,Angle,Circ.,Feret,Median,Skew,Kurt,FeretX,FeretY,FeretAngle,MinFeret,AR,Round,Solidity
#########################################################################################################################
    # Calcular la nueva columna AverageFeret (AVFeret) como el promedio de 'Feret' y 'MinFeret'
    # Calcular la nueva columna GD (Geometric Diameter) como 'Perim.' / π
    # Calcular la nueva columna AED (Area Equivalent Diameter) como 2 * √(Area / π)
# Definición de nombres descriptivos para los parámetros

# Definición de nombres descriptivos para los parámetros

# Diccionario de parámetros y nombres a mostrar
param_names = {
    'Feret': 'MaxFeret (nm)',
    'MinFeret': 'MinFeret (nm)',
    'Area': 'Area (nm²)',  # Incluye unidades
    'GD': 'GD (nm)', # Geometric Diameter 
    'Circ.': 'Circularity',
    'Round': 'Roundness',
    'AED': 'AED (nm)', # Area Equivalent Diameter
    'Perim.': 'Perimeter (nm)',
    'AvFeret': 'Average Feret (nm)',
    'Average': 'Average',
    'Median': 'Median',
    'Mode': 'Mode',
    'Std Dev': 'StdDev', # Standard Deviation
    'Q1': 'Q1', # 1st Quartile 
    'Q3': 'Q3'  # 3rd Quartile 
}

# Lista de columnas/parametros para los que queremos calcular las estadísticas
params = ['Feret', 'AvFeret', 'MinFeret', 'Area', 'GD', 'AED', 'Circ.', 'Round']


# %%%%%%%%% TABLES %%
#
# Lista de columnas para las que queremos calcular las estadísticas
columns = list(param_names.keys())  # Usar las claves del diccionario

# Diccionario para almacenar los DataFrames de estadísticas por muestra
stats_dfs = {}

# Calcular estadísticas para cada muestra
for nsample in range(1, ns + 1):
    sample_df = df_combined[df_combined['Sample'] == f'Sample {nsample}']
    
    # Verificar que las columnas existen en sample_df
    existing_columns = sample_df.columns.intersection(columns)

    stats_dict = {'Average': [], 'Median': [], 'Mode': [], 'Std Dev': [], 'Q1': [], 'Q3': []}
    
    for col in existing_columns:
        stats_dict['Average'].append(round(sample_df[col].mean(), 2))
        stats_dict['Median'].append(round(sample_df[col].median(), 2))
        stats_dict['Mode'].append(round(stats.mode(sample_df[col])[0][0], 2))
        stats_dict['Std Dev'].append(round(sample_df[col].std(), 2))
        stats_dict['Q1'].append(round(sample_df[col].quantile(0.25), 2))
        stats_dict['Q3'].append(round(sample_df[col].quantile(0.75), 2))

    # Convertir el diccionario en un DataFrame y almacenarlo en el diccionario de DataFrames
    stats_dfs[f'Sample {nsample}'] = pd.DataFrame(stats_dict, index=existing_columns)

# Mostrar las tablas de estadísticas para cada muestra
for sample, stats_df in stats_dfs.items():
    print(f"\nEstadísticas para {sample}:\n", stats_df)



# %%%%%%%%% PLOTS %%
# %%%%%%%%% TABLES WORD _ FORMATED%%

# Función para establecer el color de fondo de una celda
def set_cell_color(cell, color):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

# Función para crear la tabla en Word
def create_word_table(df, sample_name, document):
    # Crear título con el nombre de la muestra
    title = document.add_paragraph(f"{sample_name}: Statistical descriptors Table")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.runs[0]
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(31, 73, 125)  # Azul oscuro
    
    # Crear la tabla
    table = document.add_table(rows=len(params) + 1, cols=df.shape[1] + 1)  # +1 para la columna de parámetros
    table.alignment = WD_TABLE_ALIGNMENT.CENTER  # Alinea la tabla al centro del documento

    # Agregar encabezado con formato (negrita, color azul)
    table.cell(0, 0).text = 'Parameter'
    header_cell = table.cell(0, 0)
    header_cell.paragraphs[0].runs[0].font.bold = True
    header_cell.paragraphs[0].runs[0].font.size = Pt(12)
    header_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    header_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar el texto en la primera fila
    set_cell_color(header_cell, '4472C4')  # Azul oscuro

    # Llenar el encabezado para las columnas de estadísticas
    for j, col in enumerate(df.columns):
        header_cell = table.cell(0, j + 1)
        descriptive_name = param_names.get(col, col)  # Usar el nombre original si no está en el diccionario
        header_cell.text = descriptive_name
        header_cell.paragraphs[0].runs[0].font.bold = True
        header_cell.paragraphs[0].runs[0].font.size = Pt(12)
        header_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        header_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar el texto en la primera fila
        set_cell_color(header_cell, '4472C4')  # Azul oscuro

    # Agregar filas de datos con los parámetros en el orden de la lista 'params'
    for i, param in enumerate(params):
        if param not in df.index:
            continue  # Saltar si el parámetro no está en el DataFrame

        # Alternar color de fondo en filas
        row_color = 'D9E1F2' if i % 2 == 0 else 'FFFFFF'  # Gris claro alternado

        param_cell = table.cell(i + 1, 0)
        param_cell.text = param_names.get(param, param)
        param_cell.paragraphs[0].runs[0].font.bold = True
        set_cell_color(param_cell, row_color)

        for j, col in enumerate(df.columns):
            cell = table.cell(i + 1, j + 1)
            cell.text = str(df.loc[param, col])
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar los valores
            set_cell_color(cell, row_color)
    
    # Añadir un salto de línea después de cada tabla
    document.add_paragraph("")

# Crear un solo documento para todas las tablas de estadísticas
output_file_name = "combined_stats.docx"
document = Document()

# Crear tablas para cada muestra
for sample, stats_df in stats_dfs.items():
    sample_name = sample  # El nombre de la muestra (ejemplo: "Sample 1")
    create_word_table(stats_df, sample_name, document)

# Guardar el documento con todas las tablas
document.save(output_file_name)

print(f"Documento creado: {output_file_name}")
